#!/usr/bin/env python
"""Create a DOCX document from a compact JSON content description."""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path


def load_docx_module():
    try:
        import docx  # type: ignore
        from docx.enum.text import WD_BREAK  # type: ignore
        from docx.shared import Inches, Pt  # type: ignore
    except Exception as exc:
        raise RuntimeError(
            'python-docx is required. Install it with the resolved Python executable: "%FLAZZ_PYTHON%" -m pip install python-docx.'
        ) from exc
    return docx, WD_BREAK, Inches, Pt


def add_item(document, item, WD_BREAK, Inches) -> None:
    kind = item.get("type", "paragraph")
    if kind == "paragraph":
        document.add_paragraph(str(item.get("text", "")))
    elif kind == "heading":
        document.add_heading(str(item.get("text", "")), level=int(item.get("level", 1)))
    elif kind == "list":
        style = "List Number" if item.get("ordered") else "List Bullet"
        for text in item.get("items", []):
            document.add_paragraph(str(text), style=style)
    elif kind == "table":
        headers = [str(x) for x in item.get("headers", [])]
        rows = [[str(cell) for cell in row] for row in item.get("rows", [])]
        col_count = max(len(headers), *(len(row) for row in rows), 1)
        table = document.add_table(rows=1 if headers else 0, cols=col_count)
        table.style = item.get("style", "Table Grid")
        if headers:
            for idx, value in enumerate(headers):
                table.rows[0].cells[idx].text = value
        for row in rows:
            cells = table.add_row().cells
            for idx, value in enumerate(row):
                cells[idx].text = value
    elif kind == "image":
        path = item.get("path")
        if not path:
            return
        width = item.get("width_inches")
        document.add_picture(path, width=Inches(float(width)) if width else None)
    elif kind == "pageBreak":
        paragraph = document.add_paragraph()
        paragraph.add_run().add_break(WD_BREAK.PAGE)
    else:
        document.add_paragraph(str(item.get("text", "")))


def create(config: dict, output: Path) -> None:
    docx, WD_BREAK, Inches, Pt = load_docx_module()
    document = docx.Document()

    style = document.styles["Normal"]
    style.font.name = config.get("font", "Calibri")
    style.font.size = Pt(float(config.get("font_size", 11)))

    title = config.get("title")
    if title:
        document.add_heading(str(title), level=0)
    subtitle = config.get("subtitle")
    if subtitle:
        document.add_paragraph(str(subtitle))
    author = config.get("author")
    if author:
        document.add_paragraph(str(author))

    for section in config.get("sections", []):
        heading = section.get("heading")
        if heading:
            document.add_heading(str(heading), level=int(section.get("level", 1)))
        for item in section.get("content", []):
            add_item(document, item, WD_BREAK, Inches)

    for item in config.get("content", []):
        add_item(document, item, WD_BREAK, Inches)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description="Create a DOCX from JSON content.")
    parser.add_argument("--content-json", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()

    try:
        config = json.loads(args.content_json.read_text(encoding="utf-8"))
        create(config, args.output)
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1

    print(f"Created {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
